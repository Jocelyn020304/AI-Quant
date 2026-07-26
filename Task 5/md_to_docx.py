"""将 TASK5 Markdown 报告转换为格式规范的 Word 文档（最终版）
公式采用「OMML 原生公式」方案：
  LaTeX → MathML（latex2mathml）→ OMML（Python 递归构建）
  插入为 Word 原生 m:oMath 对象（与 Word 公式编辑器效果一致）

格式：宋体五号(10.5pt)、1.5倍行距、两端对齐、A4纸、页边距2.5cm
正文无加粗（仅标题加粗）、无多余星号、无乱码、无反斜杠泄漏。
"""
import re
import os
from lxml import etree

from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ═══════════════════════════════════════════════════════
#  命名空间常量
# ═══════════════════════════════════════════════════════

MATHML_NS = 'http://www.w3.org/1998/Math/MathML'
OMML_NS = 'http://schemas.openxmlformats.org/officeDocument/2006/math'

MML = '{%s}' % MATHML_NS       # MathML 命名空间前缀
M = '{%s}' % OMML_NS           # OMML 命名_space前缀

# ═══════════════════════════════════════════════════════
#  MathML → OMML 直接递归转换器（Python 实现）
# ═══════════════════════════════════════════════════════


def _make_element(tag, **attrs):
    """创建带命名空间的 OMML 元素"""
    elem = etree.SubElement(etree.Element('dummy'), M + tag)
    for k, v in attrs.items():
        elem.set(M + k, str(v))
    # 返回深拷贝（脱离 dummy 父节点）
    copy = deepcopy(elem)
    return copy


def _make_run(text):
    """创建 m:r > m:t 文本运行（自动过滤 Word 不支持的罕见 Unicode 字符）"""
    text = _sanitize_math_text(str(text))
    r = etree.Element(M + 'r')
    t = etree.SubElement(r, M + 't')
    t.text = text
    # 设置 xml:space="preserve" 以保留空格
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    return r


def _sanitize_math_text(text):
    """轻量级过滤器：仅丢弃组合变音符号（U+0300-U+036F）和控制字符。"""
    result = []
    for ch in text:
        cp = ord(ch)
        if 0x0300 <= cp <= 0x036F:
            continue
        if cp <= 0x001F or (0x007F <= cp <= 0x009F):
            continue
        result.append(ch)
    return ''.join(result)


def _map_math_alpha(ch):
    """将数学字母数字符号区块（U+1D400–U+1D7FF）映射回基础拉丁/希腊字母"""
    cp = ord(ch)

    # 粗体大写 A-Z: U+1D400–U+1D419
    if 0x1D400 <= cp <= 0x1D419:
        return chr(cp - 0x1D400 + ord('A'))
    # 粗体小写 a-z: U+1D41A–U+1D433
    if 0x1D41A <= cp <= 0x1D433:
        return chr(cp - 0x1D41A + ord('a'))
    # 粗体斜体大写: U+1D434–U+1D44D
    if 0x1D434 <= cp <= 0x1D44D:
        return chr(cp - 0x1D434 + ord('A'))
    # 粗体斜体小写: U+1D44E–U+1D467
    if 0x1D44E <= cp <= 0x1D467:
        return chr(cp - 0x1D44E + ord('a'))

    # 粗体希腊（常见范围，不完整但覆盖主要字母）
    # 𝚨 𝛃 𝛄 etc.
    # 小写粗体希腊 α-ω: 约 U+1D6C2–U+1D6DB
    if 0x1D6C2 <= cp <= 0x1D6DB:
        greek_lower = 'αβγδεζηθικλμνξοπρστυφχψω'
        idx = cp - 0x1D6C2
        if idx < len(greek_lower):
            return greek_lower[idx]

    # 大写粗体希腊 Α-Ω: 约 U+1D6A8–U+1D6C1
    if 0x1D6A8 <= cp <= 0x1D6C1:
        greek_upper = 'ΑΒΓΔΕΖΗΘΙΚΛΜΝΞΟΠΡΣΤΥΦΧΨΩ'
        idx = cp - 0x1D6A8
        if idx < len(greek_upper):
            return greek_upper[idx]

    # 粗体斜体希腊小写: 约 U+1D6FC–U+1D715
    if 0x1D6FC <= cp <= 0x1D715:
        greek_lower = 'αβγδεζηθικλμνξοπρστυφχψω'
        idx = cp - 0x1D6FC
        if idx < len(greek_lower):
            return greek_lower[idx]

    # 空集 ∅ 等特殊符号
    special_map = {
        0x1D7D8: '0', 0x1D7D9: '1', 0x1D7DA: '2',
        0x1D7DB: '3', 0x1D7DC: '4', 0x1D7DD: '5',
        0x1D7DE: '6', 0x1D7DF: '7', 0x1D7E0: '8', 0x1D7E1: '9',
    }
    if cp in special_map:
        return special_map[cp]

    return None  # 无法映射


def _make_delimited(children, beg_chr='', end_chr=''):
    """创建 m:d 分隔组（用于 mrow 等）"""
    d = etree.Element(M + 'd')
    dPr = etree.SubElement(d, M + 'dPr')
    if beg_chr or end_chr:
        beg = etree.SubElement(dPr, M + 'begChr')
        beg.set(M + 'val', beg_chr)
        end = etree.SubElement(dPr, M + 'endChr')
        end.set(M + 'val', end_chr)
    e_elem = etree.SubElement(d, M + 'e')
    for child in children:
        e_elem.append(deepcopy(child))
    return d


def deepcopy(elem):
    """深拷贝 XML 元素"""
    from copy import deepcopy as dc
    return dc(elem)


def convert_mathml_to_omml(mathml_elem):
    """
    将单个 MathML 元素递归转换为一个或多个 OMML 子元素。
    返回 OMML 元素列表。
    """
    tag = mathml_elem.tag
    local_name = tag.replace(MML, '') if tag.startswith(MML) else tag

    # ── 文本节点 ──
    if not isinstance(tag, str):  # 文本/注释/PI
        return [_make_run(mathml_elem)]

    # ── <mi> 数学变量（斜体）──
    if local_name == 'mi':
        text = (mathml_elem.text or '').strip()
        if not text:
            return []
        r = _make_run(text)
        # 检查 mathvariant 属性
        variant = mathml_elem.get('mathvariant', '')
        if variant == 'bold':
            sty = etree.Element(M + 'rPr')
            b = etree.SubElement(sty, M + 'b')
            r.insert(0, sty)
        elif variant in ('bold-italic', 'bold italic'):
            sty = etree.Element(M + 'rPr')
            b = etree.SubElement(sty, M + 'b')
            i = etree.SubElement(sty, M + 'i')
            r.insert(0, sty)
        return [r]

    # ── <mn> 数字 ──
    if local_name == 'mn':
        text = (mathml_elem.text or '').strip()
        return [_make_run(text)] if text else []

    # ── <mo> 运算符 ──
    if local_name == 'mo':
        text = (mathml_elem.text or '').strip()
        if not text:
            return []
        r = _make_run(text)
        stretchy = mathml_elem.get('stretchy', 'false')
        if stretchy == 'true':
            # 大型运算符（括号等），可能需要特殊处理
            pass
        return [r]

    # ── <mtext> 文本 ──
    if local_name == 'mtext':
        text = (mathml_elem.text or '').strip().replace('\u00a0', ' ')
        return [_make_run(text)] if text else []

    # ── <mspace> 空格 ──
    if local_name == 'mspace':
        width = mathml_elem.get('width', '0.5em')
        return [_make_run(' ')]

    # ── <mrow> 行/组 ──
    # mrow 只是逻辑分组，OMML 不需要包装元素，直接平铺子元素即可。
    # 注意：绝不能用 m:d 包装——OMML 的 m:d 默认渲染为圆括号 ()！
    if local_name == 'mrow':
        children = []
        for child in mathml_elem:
            children.extend(convert_mathml_to_omml(child))
        return children

    # ── <mfrac> 分式 ──
    if local_name == 'mfrac':
        sub_elems = list(mathml_elem)
        num_elems = convert_mathml_to_omml(sub_elems[0]) if len(sub_elems) > 0 else [_make_run('?')]
        den_elems = convert_mathml_to_omml(sub_elems[1]) if len(sub_elems) > 1 else [_make_run('?')]
        frac = etree.Element(M + 'f')
        num_node = etree.SubElement(frac, M + 'num')
        for ne in num_elems:
            num_node.append(deepcopy(ne))
        den_node = etree.SubElement(frac, M + 'den')
        for de in den_elems:
            den_node.append(deepcopy(de))
        return [frac]

    # ── <msup> 上标 ──
    if local_name == 'msup':
        sub_elems = list(mathml_elem)
        base_elem = sub_elems[0] if len(sub_elems) > 0 else None
        LARGE_OPERATORS = {'\u2211', '\u220F', '\u222B', '\u222C', '\u222D',
                            '\u22C0', '\u22C1', '\u22C2', '\u22C3', '\u2210',
                            '\u2A00', '\u2A01', '\u2A02'}
        base_text = (base_elem.text or '').strip() if base_elem is not None else ''
        if not base_text and base_elem is not None and base_elem.tag == MML + 'mo':
            base_text = (base_elem.text or '').strip()

        if base_text in LARGE_OPERATORS:
            nary = etree.Element(M + 'nary')
            naryPr = etree.SubElement(nary, M + 'naryPr')
            chr_elem = etree.SubElement(naryPr, M + 'chr')
            chr_elem.set(M + 'val', base_text)
            limLoc = etree.SubElement(naryPr, M + 'limLoc')
            limLoc.set(M + 'val', 'undOvr')
            if len(sub_elems) >= 2:
                sup_node = etree.SubElement(nary, M + 'sup')
                sup_omml = convert_mathml_to_omml(sub_elems[1])
                for spo in sup_omml:
                    sup_node.append(deepcopy(spo))
            e_node = etree.SubElement(nary, M + 'e')
            e_r = etree.SubElement(e_node, M + 'r'); e_t = etree.SubElement(e_r, M + 't'); e_t.text = ''
            return [nary]

        # 普通上下标（注意：lxml Element 无子元素时布尔值为 False，必须用 is not None 判断）
        base = convert_mathml_to_omml(base_elem) if base_elem is not None else [_make_run('?')]
        sup = convert_mathml_to_omml(sub_elems[1]) if len(sub_elems) > 1 else [_make_run('?')]
        sSup = etree.Element(M + 'sSup')
        e_node = etree.SubElement(sSup, M + 'e')
        for b in base:
            e_node.append(deepcopy(b))
        sup_node = etree.SubElement(sSup, M + 'sup')
        for s in sup:
            sup_node.append(deepcopy(s))
        return [sSup]

    # ── <msub> 下标 ──
    if local_name == 'msub':
        sub_elems = list(mathml_elem)
        base_elem = sub_elems[0] if len(sub_elems) > 0 else None
        LARGE_OPERATORS = {'\u2211', '\u220F', '\u222B', '\u222C', '\u222D',
                            '\u22C0', '\u22C1', '\u22C2', '\u22C3', '\u2210'}
        base_text = (base_elem.text or '').strip() if base_elem is not None else ''
        if not base_text and base_elem is not None and base_elem.tag == MML + 'mo':
            base_text = (base_elem.text or '').strip()

        if base_text in LARGE_OPERATORS:
            nary = etree.Element(M + 'nary')
            naryPr = etree.SubElement(nary, M + 'naryPr')
            chr_elem = etree.SubElement(naryPr, M + 'chr'); chr_elem.set(M + 'val', base_text)
            limLoc = etree.SubElement(naryPr, M + 'limLoc'); limLoc.set(M + 'val', 'undOvr')
            if len(sub_elems) >= 2:
                sub_node = etree.SubElement(nary, M + 'sub')
                sub_omml = convert_mathml_to_omml(sub_elems[1])
                for so in sub_omml: sub_node.append(deepcopy(so))
            e_node = etree.SubElement(nary, M + 'e')
            e_r = etree.SubElement(e_node, M + 'r'); e_t = etree.SubElement(e_r, M + 't'); e_t.text = ''
            return [nary]

        base = convert_mathml_to_omml(base_elem) if base_elem is not None else [_make_run('?')]
        sub = convert_mathml_to_omml(sub_elems[1]) if len(sub_elems) > 1 else [_make_run('?')]
        sSub = etree.Element(M + 'sSub')
        e_node = etree.SubElement(sSub, M + 'e')
        for b in base:
            e_node.append(deepcopy(b))
        sub_node = etree.SubElement(sSub, M + 'sub')
        for s in sub:
            sub_node.append(deepcopy(s))
        return [sSub]

    # ── <msubsup> 上下标 ──
    if local_name == 'msubsup':
        sub_elems = list(mathml_elem)
        base_elem = sub_elems[0] if len(sub_elems) > 0 else None

        # ═══ 检测大型运算符（∑ ∏ ∫ ⋂ ⋃ 等）作为底数 ═══
        # latex2mathml 对 \sum 产生 <msubsup><mo>∑</mo>...</msubsup>
        # 而非 <munderover>，需特殊处理为 OMML m:nary
        LARGE_OPERATORS = {'\u2211', '\u220F', '\u222B', '\u222C', '\u222D',
                            '\u22C0', '\u22C1', '\u22C2', '\u22C3', '\u2210',
                            '\u2A00', '\u2A01', '\u2A02', '\u2A03', '\u2A04',
                            '\u2A05', '\u2A06', '\u2A07', '\u2A08', '\u2A09',
                            '∑', '∏', '∫'}
        base_text = ''
        if base_elem is not None:
            base_text = (base_elem.text or '').strip()
            # 也检查 <mo> 内的直接文本
            if not base_text and base_elem.tag == MML + 'mo':
                base_text = (base_elem.text or '').strip()

        if base_text in LARGE_OPERATORS:
            nary = etree.Element(M + 'nary')
            naryPr = etree.SubElement(nary, M + 'naryPr')
            chr_elem = etree.SubElement(naryPr, M + 'chr')
            chr_elem.set(M + 'val', base_text)
            limLoc = etree.SubElement(naryPr, M + 'limLoc')
            limLoc.set(M + 'val', 'undOvr')

            # 下标
            if len(sub_elems) >= 2:
                sub_node = etree.SubElement(nary, M + 'sub')
                sub_omml = convert_mathml_to_omml(sub_elems[1])
                for so in sub_omml:
                    sub_node.append(deepcopy(so))

            # 上标
            if len(sub_elems) >= 3:
                sup_node = etree.SubElement(nary, M + 'sup')
                sup_omml = convert_mathml_to_omml(sub_elems[2])
                for spo in sup_omml:
                    sup_node.append(deepcopy(spo))

            # 空底数
            e_node = etree.SubElement(nary, M + 'e')
            e_r = etree.SubElement(e_node, M + 'r')
            e_t = etree.SubElement(e_r, M + 't')
            e_t.text = ''
            return [nary]

        # 普通上下标
        base = convert_mathml_to_omml(base_elem) if base_elem is not None else [_make_run('?')]
        sub = convert_mathml_to_omml(sub_elems[1]) if len(sub_elems) > 1 else [_make_run('?')]
        sup = convert_mathml_to_omml(sub_elems[2]) if len(sub_elems) > 2 else [_make_run('?')]
        sSubSup = etree.Element(M + 'sSubSup')
        e_node = etree.SubElement(sSubSup, M + 'e')
        for b in base:
            e_node.append(deepcopy(b))
        sub_node = etree.SubElement(sSubSup, M + 'sub')
        for s in sub:
            sub_node.append(deepcopy(s))
        sup_node = etree.SubElement(sSubSup, M + 'sup')
        for sp in sup:
            sup_node.append(deepcopy(sp))
        return [sSubSup]

    # ── <munderover> / <munder> / <mover> 大型运算符（求和/积分等）或变音符号 ──
    if local_name in ('munderover', 'munder', 'mover'):
        sub_elems = list(mathml_elem)

        # ═══ 检测是否为变音符号（accent） ═══
        # \tilde \hat \bar \vec \dot \breve 等产生 <mover>（有时带 accent="true"）
        # 特征：恰好2个子元素，第二个是单个运算符字符
        is_accent = False
        acc_char = ''
        ACCENT_CHARS = {'~', '^', '¯', '→', '˙', 'ˇ', '`', '´', '¨', '°'}
        if len(sub_elems) == 2 and local_name == 'mover':
            second = sub_elems[1]
            second_text = ''
            if second.tag == MML + 'mo' and (second.text or '').strip():
                second_text = (second.text or '').strip()
            elif second.tag == MML + 'mrow':
                inner_mos = [c.text or '' for c in second if c.tag == MML + 'mo']
                if len(inner_mos) == 1:
                    second_text = inner_mos[0].strip()
            if second_text and (second_text in ACCENT_CHARS or ord(second_text[0]) > 0x2F0):
                is_accent = True
                acc_char = second_text

        if is_accent and len(sub_elems) >= 1:
            # 映射到 OMML m:acc 元素
            base_elems = convert_mathml_to_omml(sub_elems[0])
            acc = etree.Element(M + 'acc')
            accPr = etree.SubElement(acc, M + 'accPr')
            chr_elem = etree.SubElement(accPr, M + 'chr')
            # 映射常见变音符号到 OMML 兼容字符
            acc_char_map = {
                '~': '\u0303',   # combining tilde → 用 ~ 本身
                '^': '^',
                '¯': '¯',
                '→': '→',
                '˙': '\u0307',   # combining dot
                'ˇ': '\u02C7',   # caron
                '`': '`',
                '´': '´',
                '¨': '¨',
                '°': '°',
            }
            omml_acc = acc_char_map.get(acc_char, acc_char)
            chr_elem.set(M + 'val', omml_acc)
            e_node = etree.SubElement(acc, M + 'e')
            for be in base_elems:
                e_node.append(deepcopy(be))
            return [acc]

        # ═══ 否则：大型运算符（求和/积分等） ═══
        sub_elems = list(mathml_elem)

        # 第一个子元素是运算符本身
        op_elem = sub_elems[0] if len(sub_elems) > 0 else None
        op_text = ''
        if op_elem is not None:
            op_children = convert_mathml_to_omml(op_elem)
            # 取运算符文本
            if op_children and op_children[0].find(M + 't') is not None:
                op_text = op_children[0].find(M + 't').text or ''

        # 判断运算符类型并映射到 OMML 字符
        char_map = {
            '\u2211': '\u2211',   # ∑ sum
            '\u220f': '\u220f',   # ∏ prod
            '\u222b': '\u222b',   # int
        }
        omml_char = char_map.get(op_text, op_text or '?')

        nary = etree.Element(M + 'nary')
        naryPr = etree.SubElement(nary, M + 'naryPr')
        chr_elem = etree.SubElement(naryPr, M + 'chr')
        chr_elem.set(M + 'val', omml_char)
        limLoc = etree.SubElement(naryPr, M + 'limLoc')
        limLoc.set(M + 'val', 'undOvr')

        # 下标（第2个子元素）
        if local_name in ('munder', 'munderover') and len(sub_elems) >= 2:
            sub_node = etree.SubElement(nary, M + 'sub')
            sub_omml = convert_mathml_to_omml(sub_elems[1])
            for so in sub_omml:
                sub_node.append(deepcopy(so))

        # 上标（第3个子元素）
        if local_name in ('mover', 'munderover') and len(sub_elems) >= 3:
            sup_node = etree.SubElement(nary, M + 'sup')
            sup_omml = convert_mathml_to_omml(sub_elems[2])
            for spo in sup_omml:
                sup_node.append(deepcopy(spo))

        # 底数（运算符后面的内容——对大型运算符通常没有底数，用空占位）
        e_node = etree.SubElement(nary, M + 'e')
        e_r = etree.SubElement(e_node, M + 'r')
        e_t = etree.SubElement(e_r, M + 't')
        e_t.text = ''  # 空底数

        return [nary]

    # ── <mroot> 方根 ──
    if local_name == 'mroot':
        sub_elems = list(mathml_elem)
        radicand = convert_mathml_to_omml(sub_elems[0]) if len(sub_elems) > 0 else [_make_run('?')]
        degree = convert_mathml_to_omml(sub_elems[1]) if len(sub_elems) > 1 else [_make_run('2')]
        rad = etree.Element(M + 'rad')
        radPr = etree.SubElement(rad, M + 'radPr')  # 默认根号样式
        deg_node = etree.SubElement(rad, M + 'deg')
        for de in degree:
            deg_node.append(deepcopy(de))
        e_node = etree.SubElement(rad, M + 'e')
        for ra in radicand:
            e_node.append(deepcopy(ra))
        return [rad]

    # ── <mfenced> 带括号组 ──
    if local_name == 'mfenced':
        open_chr = mathml_elem.get('open', '(')
        close_chr = mathml_elem.get('close', ')')
        children = []
        for child in mathml_elem:
            children.extend(convert_mathml_to_omml(child))
        if not children:
            children = [_make_run('')]
        return [_make_delimited(children, beg_chr=open_chr, end_chr=close_chr)]

    # ── <mtable> 矩阵/表格 ──
    if local_name == 'mtable':
        matrix = etree.Element(M + 'm')
        mr = etree.SubElement(matrix, M + 'mr')
        for row in mathml_elem:
            if row.tag.replace(MML, '') == 'mtr':
                e_node = etree.SubElement(mr, M + 'e')
                cell_contents = []
                for cell in row:
                    if cell.tag.replace(MML, '') in ('mtd', 'mrow'):
                        cell_contents.extend(convert_mathml_to_omml(cell))
                if not cell_contents:
                    cell_contents = [_make_run('')]
                for cc in cell_contents:
                    e_node.append(deepcopy(cc))
        return [matrix]

    # ── <mtr> / <mtd> （在 mtable 中已处理）──
    if local_name in ('mtr', 'mtd'):
        children = []
        for child in mathml_elem:
            children.extend(convert_mathml_to_omml(child))
        return children if children else [_make_run('')]

    # ── <mpadded> / <mstyle> / <semantics> / <annotation> ──
    if local_name in ('mpadded', 'mstyle', 'semantics', 'annotation-xml'):
        children = []
        for child in mathml_elem:
            children.extend(convert_mathml_to_omml(child))
        return children if children else []

    # ── 默认：递归子节点 ──
    children = []
    for child in mathml_elem:
        children.extend(convert_mathml_to_omml(child))
    return children if children else []


def latex_to_omml(latex_str):
    """
    将 LaTeX 公式字符串转为 OMML oMath 元素。
    流程：LaTeX字符串 → latex2mathml → MathML XML → 递归转换 → OMML
    """
    import latex2mathml.converter as l2m

    try:
        mml_str = l2m.convert(latex_str)
    except Exception as e:
        # 回退：纯文本包装
        mml_str = (
            '<math xmlns="http://www.w3.org/1998/Math/MathML">'
            '<mrow><mi>' + _xml_escape(latex_str) + '</mi></mrow></math>'
        )

    parser = etree.XMLParser(remove_blank_text=True)
    try:
        mathml_root = etree.fromstring(mml_str.encode('utf-8'), parser=parser)
    except Exception:
        mathml_root = etree.fromstring(
            ('<math xmlns="http://www.w3.org/1998/Math/MathML">'
             '<mrow><mi>' + _xml_escape(latex_str) + '</mi></mrow></math>').encode('utf-8'),
            parser=parser
        )

    # 找到实际的数学内容（跳过 <math> 根标签）
    content = mathml_root
    if mathml_root.tag == MML + 'math' or mathml_root.tag.endswith('}math'):
        # 取第一个子元素作为内容
        children = list(mathml_root)
        if children:
            content = children[0]
        else:
            # 空的 math 标签
            content = mathml_root

    # 递归转换
    omml_children = convert_mathml_to_omml(content)

    if not omml_children:
        omml_children = [_make_run(latex_str)]

    # 包装为 oMath 元素
    oMath = etree.Element(M + 'oMath')
    for child in omml_children:
        oMath.append(deepcopy(child))

    return oMath


def _xml_escape(s):
    s = s.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
    s = s.replace('"', '&quot;').replace("'", '&apos;')
    return s


def insert_inline_omml(paragraph, omml_element):
    """将 OMML oMath 作为行内公式插入段落（不使用 oMathPara 块级容器）"""
    copied = deepcopy(omml_element)
    paragraph._element.append(copied)


def insert_display_omml(doc, latex_str):
    """将 OMML oMath 作为块级居中显示公式插入（用 oMathPara 包裹）"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(10)
    pf.space_after = Pt(10)

    omml = latex_to_omml(latex_str.strip())
    # 块级公式：oMathPara > oMath
    oMathPara = etree.SubElement(p._element, M + 'oMathPara')
    copied = deepcopy(omml)
    oMathPara.append(copied)


# ═══════════════════════════════════════════════════════
#  Word 文档构建工具函数
# ═══════════════════════════════════════════════════════

MD_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                       "TASK5_机器学习算法与金融应用.md")
OUT_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(OUT_DIR, "TASK5_机器学习算法与金融应用.docx")

doc = Document()

section = doc.sections[0]
section.page_width = Cm(21); section.page_height = Cm(29.7)
section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5); section.bottom_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = '宋体'; font.size = Pt(10.5); font.color.rgb = RGBColor(0, 0, 0)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
pf = style.paragraph_format
pf.line_spacing = 1.5; pf.space_before = Pt(0); pf.space_after = Pt(0)
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def set_run_font(run, name_cn='宋体', name_en='宋体', size=Pt(10.5),
                 bold=False, italic=False, color=RGBColor(0, 0, 0)):
    run.font.name = name_en; run.font.size = size
    run.font.bold = bold; run.font.italic = italic; run.font.color.rgb = color
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}><w:rFonts w:eastAsia="{name_cn}"/></w:rPr>')
        r.insert(0, rPr)
    else:
        eastAsia = rPr.find(qn('w:rFonts'))
        if eastAsia is None:
            eastAsia = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{name_cn}"/>')
            rPr.insert(0, eastAsia)
        else:
            eastAsia.set(qn('w:eastAsia'), name_cn)


def add_heading_styled(text, level=1):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format; pf.line_spacing = 1.5
    pf.space_before = Pt(12 if level == 1 else 6); pf.space_after = Pt(6)
    run = p.add_run(text)
    size = {1: Pt(16), 2: Pt(14), 3: Pt(12), 4: Pt(10.5)}.get(level, Pt(10.5))
    set_run_font(run, size=size, bold=True)


INLINE_RE = re.compile(r'\$([^\$]+?)\$')


def add_para(text, indent=True):
    """普通段落：$...$ 行内公式转 OMML"""
    p = doc.add_paragraph()
    pf = p.paragraph_format; pf.line_spacing = 1.5
    pf.space_before = Pt(0); pf.space_after = Pt(0); pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        pf.first_line_indent = Cm(0.74)

    clean_text = text.replace('**', '')

    last_end = 0
    for m in INLINE_RE.finditer(clean_text):
        before = clean_text[last_end:m.start()]
        if before:
            set_run_font(p.add_run(before))

        formula_latex = m.group(1).strip()
        omml = latex_to_omml(formula_latex)
        insert_inline_omml(p, omml)
        last_end = m.end()

    remaining = clean_text[last_end:]
    if remaining:
        set_run_font(p.add_run(remaining))


def set_table_width(table, width_cm=16.0):
    tbl = table._tbl; tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>'); tbl.insert(0, tblPr)
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{int(width_cm * 567)}" w:type="dxa"/>')
        tblPr.append(tblW)
    else:
        tblW.set(qn('w:w'), str(int(width_cm * 567)))


def set_cell_margins(cell, top=0, bottom=0, left=0.05, right=0.05):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{int(top*20)}" w:type="dxa"/>'
        f'<w:bottom w:w="{int(bottom*20)}" w:type="dxa"/>'
        f'<w:left w:w="{int(left*100)}" w:type="dxa"/>'
        f'<w:right w:w="{int(right*100)}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)


def add_table_from_markdown(table_lines):
    rows_data = []
    for line in table_lines:
        line = line.strip()
        if not line or re.match(r'^\|[\|\-\=\:\s]+\|$', line):
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        rows_data.append(cells)
    if not rows_data:
        return
    table = doc.add_table(rows=len(rows_data), cols=len(rows_data[0]))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.style = 'Table Grid'
    set_table_width(table, 16.0)
    for i, rd in enumerate(rows_data):
        for j, ct in enumerate(rd):
            cell = table.cell(i, j); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell); p = cell.paragraphs[0]; p.clear()
            pf = p.paragraph_format; pf.line_spacing = 1.0; pf.space_before = Pt(0); pf.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(ct.replace('**', ''))
            set_run_font(run, size=Pt(9))
            if i == 0:
                run.font.bold = True
    doc.add_paragraph()


# ═══════════════════════════════════════════════════════
#  解析 Markdown 并生成 Word
# ═══════════════════════════════════════════════════════

with open(MD_PATH, 'r', encoding='utf-8') as f:
    lines = f.readlines()

i = 0
in_math = False
math_buf = []

while i < len(lines):
    line = lines[i].rstrip('\n').rstrip('\r')
    if not line:
        i += 1; continue

    if line.startswith('$$'):
        if in_math:
            formula_text = '\n'.join(math_buf).strip()
            if formula_text:
                insert_display_omml(doc, formula_text)
            math_buf = []; in_math = False
        else:
            in_math = True
        i += 1; continue

    if in_math:
        math_buf.append(line); i += 1; continue

    h_match = re.match(r'^(#{1,4})\s+(.+)$', line)
    if h_match:
        add_heading_styled(h_match.group(2), len(h_match.group(1))); i += 1; continue

    if line.startswith('|'):
        tlines = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            tlines.append(lines[i].strip()); i += 1
        has_sep = any(re.match(r'^\|[\|\-\=\:\s]+\|$', l) for l in tlines)
        if has_sep and len(tlines) >= 2:
            add_table_from_markdown(tlines)
        else:
            for tl in tlines:
                add_para(tl.strip('|'))
        continue

    if re.match(r'^---+\s*$', line):
        doc.add_paragraph(); i += 1; continue

    img_match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
    if img_match:
        img_name = img_match.group(2)
        img_alt = img_match.group(1).strip() or '图'
        img_full = os.path.join(OUT_DIR, img_name)
        if os.path.exists(img_full):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.add_run().add_picture(img_full, width=Inches(6.0))
            cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.first_line_indent = None
            cap.paragraph_format.space_before = Pt(3); cap.paragraph_format.space_after = Pt(6)
            set_run_font(cap.add_run(img_alt), size=Pt(10.5))
        else:
            add_para(f'[图片未找到: {img_name}]')
        i += 1; continue

    add_para(line); i += 1

doc.save(OUT_PATH)
print(f"Word 文档已保存至：{OUT_PATH}")
