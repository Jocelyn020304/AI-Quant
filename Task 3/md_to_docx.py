"""将 Task 3 Markdown 转换为格式规范的 Word 文档（公式用 Unicode+Times New Roman）"""
import re
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

MD_PATH = r'C:\Users\Administrator\Desktop\量化交易：AI大模型辅助的金融交易策略\Task 3\双均线策略理论解释.md'
OUT_DIR = r'C:\Users\Administrator\Desktop\量化交易：AI大模型辅助的金融交易策略\Task 3'
OUT_PATH = r'C:\Users\Administrator\Desktop\量化交易：AI大模型辅助的金融交易策略\Task 3\TASK3_双均线策略理论分析.docx'

doc = Document()

# 页面设置
s = doc.sections[0]
s.page_width = Cm(21); s.page_height = Cm(29.7)
s.left_margin = s.right_margin = s.top_margin = s.bottom_margin = Cm(2.5)

# 全局样式
style = doc.styles['Normal']
font = style.font; font.name = '宋体'; font.size = Pt(10.5)
font.color.rgb = RGBColor(0, 0, 0)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
pf = style.paragraph_format; pf.line_spacing = 1.5
pf.space_before = Pt(0); pf.space_after = Pt(0)
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
pt0 = Pt(0)

# ─── 辅助函数 ───
def set_run_font(run, name_cn='宋体', name_en='宋体', size=Pt(10.5), bold=False, italic=False, color=RGBColor(0, 0, 0)):
    run.font.name = name_en
    run.font.size = size
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}><w:rFonts w:eastAsia="{name_cn}"/></w:rPr>')
        r.insert(0, rPr)
    else:
        ea = rPr.find(qn('w:rFonts'))
        if ea is None:
            ea = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{name_cn}"/>')
            rPr.insert(0, ea)
        else:
            ea.set(qn('w:eastAsia'), name_cn)

def add_heading_styled(text, level=1):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format; pf.line_spacing = 1.5
    pf.space_before = Pt(12 if level == 1 else 6); pf.space_after = Pt(6)
    run = p.add_run(text)
    sz = {1: Pt(16), 2: Pt(14), 3: Pt(12), 4: Pt(10.5)}.get(level, Pt(10.5))
    set_run_font(run, size=sz, bold=True)

def add_para(text, indent=True):
    p = doc.add_paragraph()
    pf = p.paragraph_format; pf.line_spacing = 1.5
    pf.space_before = pt0; pf.space_after = pt0
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent: pf.first_line_indent = Cm(0.74)
    # 处理加粗 **text** 和行内公式 $...$
    parts = re.split(r'(\*\*.*?\*\*|\$.*?\$)', text)
    for part in parts:
        if not part: continue
        if part.startswith('**') and part.endswith('**'):
            r = p.add_run(part[2:-2]); set_run_font(r, bold=True)
        elif part.startswith('$') and part.endswith('$'):
            # 行内公式：用 TNR 斜体显示
            inner = part[1:-1]
            r = p.add_run(latex_to_unicode(inner))
            set_run_font(r, name_cn='Times New Roman', name_en='Times New Roman', size=Pt(10.5), italic=True)
        else:
            r = p.add_run(part); set_run_font(r)

def set_table_width(table, wc=16.0):
    tbl = table._tbl; tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>'); tbl.insert(0, tblPr)
    tw = tblPr.find(qn('w:tblW'))
    if tw is None:
        tw = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{int(wc*567)}" w:type="dxa"/>'); tblPr.append(tw)
    else: tw.set(qn('w:w'), str(int(wc*567)))

def set_cell_margins(cell, t=0, b=0, l=0.05, r=0.05):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcPr.append(parse_xml(f'<w:tcMar {nsdecls("w")}><w:top w:w="{int(t*20)}" w:type="dxa"/><w:bottom w:w="{int(b*20)}" w:type="dxa"/><w:left w:w="{int(l*100)}" w:type="dxa"/><w:right w:w="{int(r*100)}" w:type="dxa"/></w:tcMar>'))

def add_table_from_markdown(lines):
    rows = []
    for l in lines:
        l = l.strip()
        if not l or re.match(r'^\|[\|\-=\:\s]+\|$', l): continue
        rows.append([c.strip() for c in l.split('|')[1:-1]])
    if not rows: return
    t = doc.add_table(rows=len(rows), cols=len(rows[0]))
    t.alignment = WD_TABLE_ALIGNMENT.CENTER; t.style = 'Table Grid'
    set_table_width(t, 16.0)
    for i, rd in enumerate(rows):
        for j, ct in enumerate(rd):
            c = t.cell(i, j); c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(c); p = c.paragraphs[0]; p.clear()
            p.paragraph_format.line_spacing = 1.0; p.paragraph_format.space_before = pt0
            p.paragraph_format.space_after = pt0; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(ct); set_run_font(run, size=Pt(9))
            if i == 0: run.font.bold = True
    doc.add_paragraph()

# ─── LaTeX 块级公式 → Unicode 排版 ───
GREEK = {'alpha':'α','beta':'β','sigma':'σ','mu':'μ','delta':'δ','theta':'θ','pi':'π'}
SYMS = {'cdot':'·','times':'×','cdots':'…','ldots':'…',
        'leq':'≤','geq':'≥','le':'≤','ge':'≥','to':'→',
        'infty':'∞','partial':'∂','sum':'Σ','prod':'Π','int':'∫'}
TEXT_TAGS = ('mathrm', 'text', 'textbf', 'mathit', 'mathbf', 'operatorname')

SUB_MAP = {'0':'₀','1':'₁','2':'₂','3':'₃','4':'₄','5':'₅',
       '6':'₆','7':'₇','8':'₈','9':'₉',
       't':'ₜ','s':'ₛ','p':'ₚ','x':'ₓ','n':'ₙ','m':'ₘ',
       'k':'ₖ','j':'ⱼ','i':'ᵢ'}
SUP_MAP = {'0':'⁰','1':'¹','2':'²','3':'³','4':'⁴','5':'⁵',
       '6':'⁶','7':'⁷','8':'⁸','9':'⁹','T':'ᵀ','N':'ᴺ'}


def _find_matching_brace(s, start):
    depth = 1; i = start
    while i < len(s) and depth > 0:
        if s[i] == '{': depth += 1
        elif s[i] == '}': depth -= 1
        i += 1
    return i - 1


def _to_sub(text):
    return ''.join(SUB_MAP.get(c, c) for c in text)


def _to_sup(text):
    return ''.join(SUP_MAP.get(c, c) for c in text)


def _parse_latex(s):
    """LaTeX → 文本（支持 \\command, {}, ^, _）"""
    out = []; i = 0
    while i < len(s):
        ch = s[i]
        if ch == '\\':
            j = i + 1
            while j < len(s) and s[j].isalpha(): j += 1
            cmd = s[i+1:j]
            i = j
            if cmd == 'frac':
                # \frac{X}{Y}
                if i < len(s) and s[i] == ' ': i += 1
                num = ''
                if i < len(s) and s[i] == '{':
                    end = _find_matching_brace(s, i+1)
                    num = _parse_latex(s[i+1:end]); i = end + 1
                if i < len(s) and s[i] == '{':
                    end = _find_matching_brace(s, i+1)
                    den = _parse_latex(s[i+1:end]); i = end + 1
                else:
                    den = ''
                out.append(f'({num})/({den})')
            elif cmd in TEXT_TAGS:
                if i < len(s) and s[i] == '{':
                    end = _find_matching_brace(s, i+1)
                    out.append(s[i+1:end]); i = end + 1
            elif cmd in GREEK:
                out.append(GREEK[cmd])
            elif cmd in SYMS:
                out.append(SYMS[cmd])
            elif cmd in ('max','min','sin','cos','log','ln','exp','arg'):
                # 处理 max_{...} / min_{...}
                if i < len(s) and s[i] == '_':
                    i += 1
                    if i < len(s) and s[i] == '{':
                        end = _find_matching_brace(s, i+1)
                        inner = _parse_latex(s[i+1:end])
                        out.append(f'{cmd}_{inner}')
                        i = end + 1
                    elif i < len(s):
                        out.append(f'{cmd}_' + _to_sub(s[i]))
                        i += 1
                else:
                    out.append(cmd)
            else:
                out.append('\\' + cmd)
        elif ch == '^':
            i += 1
            if i < len(s) and s[i] == '{':
                end = _find_matching_brace(s, i+1)
                inner = _parse_latex(s[i+1:end]); i = end + 1
                out.append(_to_sup(inner))
            elif i < len(s):
                out.append(_to_sup(s[i])); i += 1
        elif ch == '_':
            i += 1
            if i < len(s) and s[i] == '{':
                end = _find_matching_brace(s, i+1)
                inner = _parse_latex(s[i+1:end]); i = end + 1
                out.append(_to_sub(inner))
            elif i < len(s):
                out.append(_to_sub(s[i])); i += 1
        elif ch == '{':
            end = _find_matching_brace(s, i+1)
            out.append(_parse_latex(s[i+1:end])); i = end + 1
        elif ch == ' ':
            if out and not out[-1].endswith(' '):
                out.append(' ')
            i += 1
        else:
            out.append(ch); i += 1
    return ''.join(out)


def latex_to_unicode(latex):
    return _parse_latex(latex.strip())

    # 处理下标 X_{Y} → X_Y 用 unicode 下标字符（常用字母）
    SUB_MAP = {
        '0':'₀','1':'₁','2':'₂','3':'₃','4':'₄','5':'₅',
        '6':'₆','7':'₇','8':'₈','9':'₉',
        't':'ₜ','s':'ₛ','p':'ₚ','f':'f','c':'c','x':'ₓ',
        'n':'ₙ','m':'ₘ','k':'ₖ','T':'T','N':'N','0':'₀',
    }
    def repl_sub(m):
        b = m.group(1)
        sub = m.group(2)
        sub_chars = ''.join(SUB_MAP.get(ch, ch) for ch in sub)
        return f'{b}{sub_chars}'
    s = re.sub(r'([A-Za-zα-ω0-9\\]+)_\{([^{}]*)\}', repl_sub, s)
    s = re.sub(r'([A-Za-zα-ω0-9\\])_([A-Za-z0-9])', lambda m: f'{m.group(1)}{SUB_MAP.get(m.group(2), m.group(2))}', s)

    # 处理上标 X^{Y} → 用 unicode 上标
    SUP_MAP = {
        '0':'⁰','1':'¹','2':'²','3':'³','4':'⁴','5':'⁵',
        '6':'⁶','7':'⁷','8':'⁸','9':'⁹',
        'T':'ᵀ','N':'ᴺ',
    }
    def repl_sup(m):
        b = m.group(1)
        sup = m.group(2)
        sup_chars = ''.join(SUP_MAP.get(ch, ch) for ch in sup)
        return f'{b}{sup_chars}'
    s = re.sub(r'([A-Za-zα-ω0-9\\]+)\^\{([^{}]*)\}', repl_sup, s)

    # 希腊字母/符号
    s = s.replace('\\alpha', 'α').replace('\\beta', 'β')
    s = s.replace('\\sigma', 'σ').replace('\\mu', 'μ')
    s = s.replace('\\delta', 'δ').replace('\\theta', 'θ')
    s = s.replace('\\cdot', '·').replace('\\times', '×')
    s = s.replace('\\cdots', '…').replace('\\ldots', '…')
    s = s.replace('\\leq', '≤').replace('\\geq', '≥')
    s = s.replace('\\le', '≤').replace('\\ge', '≥')
    s = s.replace('\\to', '→')
    s = s.replace('\\max', 'max').replace('\\min', 'min')
    s = s.replace('\\sum', 'Σ').replace('\\prod', 'Π').replace('\\int', '∫')
    s = s.replace('\\infty', '∞').replace('\\partial', '∂')

    # 清理多余花括号
    s = s.replace('{', '').replace('}', '')

    # 清理多余反斜杠
    s = re.sub(r'\\([A-Za-z]+)', r'\1', s)
    return s.strip()


def add_formula_paragraph(latex):
    """居中显示数学公式，Times New Roman 字体"""
    txt = latex_to_unicode(latex)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(6)
    pf.space_after = Pt(6)
    run = p.add_run(txt)
    set_run_font(run, name_cn='Times New Roman', name_en='Times New Roman', size=Pt(11), italic=True)


# ─── 解析 Markdown ───
with open(MD_PATH, 'r', encoding='utf-8') as f:
    lines = f.readlines()

i = 0; in_math = False; math_buf = []
while i < len(lines):
    line = lines[i].rstrip()
    if not line: i += 1; continue

    if line.startswith('$$'):
        if in_math:
            add_formula_paragraph('\n'.join(math_buf).strip())
            math_buf = []; in_math = False
        else: in_math = True
        i += 1; continue
    if in_math: math_buf.append(line); i += 1; continue

    hm = re.match(r'^(#{1,4})\s+(.+)$', line)
    if hm:
        add_heading_styled(hm.group(2), len(hm.group(1)))
        i += 1; continue

    if line.startswith('|'):
        tbl = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            tbl.append(lines[i].strip()); i += 1
        if any(re.match(r'^\|[\|\-=\:\s]+\|$', l) for l in tbl) and len(tbl) >= 2:
            add_table_from_markdown(tbl)
        else:
            for tl in tbl: add_para(tl.strip('|'))
        continue

    if line.startswith('---') and len(line) >= 3: doc.add_paragraph(); i += 1; continue

    img = re.match(r'!\[(.*?)\]\((.*?)\)\s*$', line)
    if img:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(); fp = os.path.join(OUT_DIR, img.group(2))
        if os.path.exists(fp): r.add_picture(fp, width=Inches(6.0))
        else: add_para(f'[图片未找到]')
        i += 1; continue

    add_para(line, indent=True); i += 1

doc.save(OUT_PATH)
print(f"Word 文档已保存至：{OUT_PATH}")
