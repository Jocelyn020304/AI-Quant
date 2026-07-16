"""将 TASK5 Markdown 报告转换为格式规范的 Word 文档
格式：宋体五号(10.5pt)、1.5倍行距、两端对齐、A4纸、页边距2.5cm
"""
import re, os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

# ─── LaTeX 块级公式 → Unicode 排版 ───
GREEK = {'alpha':'α','beta':'β','sigma':'σ','mu':'μ','delta':'δ','theta':'θ','pi':'π',
         'Delta':'Δ','Gamma':'Γ','Sigma':'Σ','Omega':'Ω','Phi':'Φ','lambda':'λ'}
SYMS = {'cdot':'·','times':'×','cdots':'…','ldots':'…',
        'leq':'≤','geq':'≥','le':'≤','ge':'≥','to':'→','rightarrow':'→',
        'infty':'∞','partial':'∂','sum':'Σ','prod':'Π','int':'∫',
        'pm':'±','mp':'∓','nabla':'∇','approx':'≈','neq':'≠'}
TEXT_TAGS = ('mathrm', 'text', 'textbf', 'mathit', 'mathbf', 'operatorname')
NOISE = ('big', 'Big', 'bigg', 'Bigg', 'bigl', 'bigr', 'biggl', 'biggr', 'bigm',
         'left', 'right', 'quad', 'qquad', 'displaystyle', 'textstyle', 'scriptstyle')

SUB_MAP = {'0':'₀','1':'₁','2':'₂','3':'₃','4':'₄','5':'₅',
       '6':'₆','7':'₇','8':'₈','9':'₉',
       't':'ₜ','s':'ₛ','p':'ₚ','x':'ₓ','n':'ₙ','m':'ₘ',
       'k':'ₖ','j':'ⱼ','i':'ᵢ','a':'ₐ','r':'ᵣ'}
SUP_MAP = {'0':'⁰','1':'¹','2':'²','3':'³','4':'⁴','5':'⁵',
       '6':'⁶','7':'⁷','8':'⁸','9':'⁹','T':'ᵀ','N':'ᴺ','t':'ᵗ'}


def _find_matching_brace(s, start):
    depth = 1; i = start
    while i < len(s) and depth > 0:
        if s[i] == '{': depth += 1
        elif s[i] == '}': depth -= 1
        i += 1
    return i - 1

def _to_sub(text): return ''.join(SUB_MAP.get(c, c) for c in text)
def _to_sup(text): return ''.join(SUP_MAP.get(c, c) for c in text)

def _parse_latex(s):
    out = []; i = 0
    while i < len(s):
        ch = s[i]
        if ch == '\\':
            j = i + 1
            while j < len(s) and s[j].isalpha(): j += 1
            cmd = s[i+1:j]; i = j
            if cmd == 'frac':
                if i < len(s) and s[i] == ' ': i += 1
                num = ''
                if i < len(s) and s[i] == '{':
                    end = _find_matching_brace(s, i+1)
                    num = _parse_latex(s[i+1:end]); i = end + 1
                if i < len(s) and s[i] == '{':
                    end = _find_matching_brace(s, i+1)
                    den = _parse_latex(s[i+1:end]); i = end + 1
                else: den = ''
                out.append(f'({num})/({den})')
            elif cmd in TEXT_TAGS:
                if i < len(s) and s[i] == '{':
                    end = _find_matching_brace(s, i+1)
                    out.append(s[i+1:end]); i = end + 1
            elif cmd in GREEK: out.append(GREEK[cmd])
            elif cmd in SYMS: out.append(SYMS[cmd])
            elif cmd in NOISE: pass
            elif cmd in ('max','min','sin','cos','log','ln','exp','lim','det','inf','sup'):
                if i < len(s) and s[i] == '_':
                    i += 1
                    if i < len(s) and s[i] == '{':
                        end = _find_matching_brace(s, i+1)
                        inner = _parse_latex(s[i+1:end]); i = end + 1
                        out.append(f'{cmd}_{inner}')
                    elif i < len(s):
                        out.append(f'{cmd}_'); i += 1
                else: out.append(cmd)
            else: pass
        elif ch == '^':
            i += 1
            if i < len(s) and s[i] == '{':
                end = _find_matching_brace(s, i+1)
                inner = _parse_latex(s[i+1:end]); i = end + 1
                out.append(_to_sup(inner))
            elif i < len(s): out.append(_to_sup(s[i])); i += 1
        elif ch == '_':
            i += 1
            if i < len(s) and s[i] == '{':
                end = _find_matching_brace(s, i+1)
                inner = _parse_latex(s[i+1:end]); i = end + 1
                out.append(_to_sub(inner))
            elif i < len(s): out.append(_to_sub(s[i])); i += 1
        elif ch == '{':
            end = _find_matching_brace(s, i+1)
            out.append(_parse_latex(s[i+1:end])); i = end + 1
        elif ch == '}': i += 1
        elif ch == ' ':
            if out and not out[-1].endswith(' '): out.append(' ')
            i += 1
        elif ch == ',': out.append(', '); i += 1
        elif ch == ';': out.append(' '); i += 1
        else: out.append(ch); i += 1
    return ''.join(out)

def latex_to_unicode(latex): return _parse_latex(latex.strip())

# ---------- 配置 ----------
MD_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), "TASK6_智能决策者机器学习策略分析.md")
OUT_DIR = os.path.dirname(os.path.abspath(__file__))
OUT_PATH = os.path.join(OUT_DIR, "TASK6_智能决策者机器学习策略分析.docx")

doc = Document()

# 页面设置 A4
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

# 全局默认样式：宋体五号，1.5倍行距，两端对齐
style = doc.styles['Normal']
font = style.font; font.name = '宋体'; font.size = Pt(10.5); font.color.rgb = RGBColor(0, 0, 0)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
pf = style.paragraph_format; pf.line_spacing = 1.5
pf.space_before = Pt(0); pf.space_after = Pt(0); pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

def set_run_font(run, name_cn='宋体', name_en='宋体', size=Pt(10.5), bold=False, italic=False, color=RGBColor(0, 0, 0)):
    run.font.name = name_en; run.font.size = size; run.font.bold = bold
    run.font.italic = italic; run.font.color.rgb = color
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}><w:rFonts w:eastAsia="{name_cn}"/></w:rPr>')
        r.insert(0, rPr)
    else:
        eastAsia = rPr.find(qn('w:rFonts'))
        if eastAsia is None:
            eastAsia = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{name_cn}"/>')
            rPr.insert(0, eastAsia)
        else: eastAsia.set(qn('w:eastAsia'), name_cn)

def add_heading_styled(text, level=1):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format; pf.line_spacing = 1.5; pf.space_before = Pt(12 if level==1 else 6); pf.space_after = Pt(6)
    run = p.add_run(text)
    size = {1:Pt(16), 2:Pt(14), 3:Pt(12), 4:Pt(10.5)}.get(level, Pt(10.5))
    set_run_font(run, size=size, bold=True)

def add_para(text, indent=True):
    p = doc.add_paragraph(); pf = p.paragraph_format; pf.line_spacing = 1.5
    pf.space_before = Pt(0); pf.space_after = Pt(0); pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent: pf.first_line_indent = Cm(0.74)
    text = text.replace('**', '')
    inline_parts = re.split(r'(\$[^\$]+\$)', text)
    for ip in inline_parts:
        if ip.startswith('$') and ip.endswith('$'):
            formula_text = latex_to_unicode(ip[1:-1])
            run = p.add_run(formula_text)
            set_run_font(run, name_cn='Times New Roman', name_en='Times New Roman', size=Pt(10.5), italic=True)
        else:
            run = p.add_run(ip); set_run_font(run)

def add_formula_paragraph(latex):
    txt = latex_to_unicode(latex)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format; pf.line_spacing = 1.5; pf.space_before = Pt(6); pf.space_after = Pt(6)
    run = p.add_run(txt)
    set_run_font(run, name_cn='Times New Roman', name_en='Times New Roman', size=Pt(11), italic=True)

def set_table_width(table, width_cm=16.0):
    tbl = table._tbl; tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None: tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>'); tbl.insert(0, tblPr)
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{int(width_cm * 567)}" w:type="dxa"/>'); tblPr.append(tblW)
    else: tblW.set(qn('w:w'), str(int(width_cm * 567)))

def set_cell_margins(cell, top=0, bottom=0, left=0.05, right=0.05):
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{int(top*20)}" w:type="dxa"/>'
        f'<w:bottom w:w="{int(bottom*20)}" w:type="dxa"/>'
        f'<w:left w:w="{int(left*100)}" w:type="dxa"/>'
        f'<w:right w:w="{int(right*100)}" w:type="dxa"/></w:tcMar>')
    tcPr.append(tcMar)

def add_table_from_markdown(table_lines):
    rows_data = []
    for line in table_lines:
        line = line.strip()
        if not line or re.match(r'^\|[\|\-\=\:\s]+\|$', line): continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        rows_data.append(cells)
    if not rows_data: return
    num_cols = len(rows_data[0])
    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER; table.style = 'Table Grid'
    set_table_width(table, 16.0)
    for i, row_data in enumerate(rows_data):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell); p = cell.paragraphs[0]; p.clear()
            pf = p.paragraph_format; pf.line_spacing = 1.0; pf.space_before = Pt(0); pf.space_after = Pt(0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER; run = p.add_run(cell_text.replace('**', '')); set_run_font(run, size=Pt(9))
            if i == 0: run.font.bold = True
    doc.add_paragraph()

# ---------- 解析 Markdown ----------
with open(MD_PATH, 'r', encoding='utf-8') as f:
    lines = f.readlines()

i = 0; in_math = False; math_buf = []

while i < len(lines):
    line = lines[i].rstrip()
    if not line: i += 1; continue

    # 数学公式块 $$...$$
    if line.startswith('$$'):
        if in_math:
            add_formula_paragraph('\n'.join(math_buf).strip()); math_buf = []; in_math = False
        else: in_math = True
        i += 1; continue
    if in_math: math_buf.append(line); i += 1; continue

    # 标题
    h_match = re.match(r'^(#{1,4})\s+(.+)$', line)
    if h_match:
        level = len(h_match.group(1)); title = h_match.group(2)
        add_heading_styled(title, level); i += 1; continue

    # 表格
    if line.startswith('|'):
        table_lines = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            table_lines.append(lines[i].strip()); i += 1
        has_separator = any(re.match(r'^\|[\|\-\=\:\s]+\|$', l) for l in table_lines)
        if has_separator and len(table_lines) >= 2: add_table_from_markdown(table_lines)
        else:
            for tl in table_lines: add_para(tl.strip('|'))
        continue

    if line.startswith('---') and len(line) >= 3: doc.add_paragraph(); i += 1; continue

    # 图片
    img_match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
    if img_match:
        img_name = img_match.group(2); img_alt = img_match.group(1).strip() or '图'
        img_path_full = os.path.join(OUT_DIR, img_name)
        if os.path.exists(img_path_full):
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER; run = p.add_run()
            run.add_picture(img_path_full, width=Inches(6.0))
            cap = doc.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.first_line_indent = None; cap.paragraph_format.space_before = Pt(3); cap.paragraph_format.space_after = Pt(6)
            cr = cap.add_run(img_alt); set_run_font(cr, size=Pt(10.5))
        else: add_para(f'[图片未找到: {img_name}]')
        i += 1; continue

    add_para(line); i += 1

# ---------- 保存 ----------
doc.save(OUT_PATH)
print(f"Word 文档已保存至：{OUT_PATH}")
