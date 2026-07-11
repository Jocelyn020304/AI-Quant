"""将 Markdown 诊断报告转换为格式规范的 Word 文档（修复版）"""
import re
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
from docx.enum.section import WD_ORIENT
import os

# ---------- 配置 ----------
MD_PATH = r'C:\Users\Administrator\Desktop\量化交易：AI大模型辅助的金融交易策略\Task 2\寒武纪_数据诊断分析报告.md'
IMG_DIR = r'C:\Users\Administrator\Desktop\量化交易：AI大模型辅助的金融交易策略\Task 2'
OUT_PATH = r'C:\Users\Administrator\Desktop\量化交易：AI大模型辅助的金融交易策略\Task 2\TASK2_寒武纪量化交易策略分析.docx'

doc = Document()


# ---------- 辅助函数：删除表格边框 ----------
def remove_table_borders(table):
    """删除表格的所有边框"""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        '<w:top w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:left w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:bottom w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:right w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideH w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '<w:insideV w:val="none" w:sz="0" w:space="0" w:color="auto"/>'
        '</w:tblBorders>'
    )
    tblPr.append(borders)

# ---------- 页面设置（A4，窄边距） ----------
section = doc.sections[0]
section.page_width = Cm(21)
section.page_height = Cm(29.7)
section.left_margin = Cm(2.5)
section.right_margin = Cm(2.5)
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)

# ---------- 全局默认样式设置 ----------
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(10.5)  # 五号
font.color.rgb = RGBColor(0, 0, 0)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

pf = style.paragraph_format
pf.line_spacing = 1.5
pf.space_before = Pt(0)
pf.space_after = Pt(0)
pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# ---------- 辅助函数：设置字体 ----------
def set_run_font(run, name_cn='宋体', name_en='宋体', size=Pt(10.5), bold=False, color=RGBColor(0, 0, 0)):
    run.font.name = name_en
    run.font.name = name_en
    run.font.size = size
    run.font.bold = bold
    run.font.color.rgb = color
    r = run._element
    rPr = r.find(qn('w:rPr'))
    if rPr is None:
        rPr = parse_xml(f'<w:rPr {nsdecls("w")}><w:rFonts w:eastAsia="{name_cn}"/></w:rPr>')
        r.insert(0, rPr)
    else:
        eastAsia = rPr.find(qn('w:rFonts'))
        if eastAsia is None:
            eastAsia = parse_xml(f'<w:rFonts {nsdecls("w")} w:eastAsia="{name_cn}"/>')
            rPr.insert(0, eastAsia)
        else:
            eastAsia.set(qn('w:eastAsia'), name_cn)

# ---------- 辅助函数：添加标题 ----------
def add_heading_styled(text, level=1):
    # 不使用 add_heading，避免默认蓝色
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(12 if level == 1 else 6)
    pf.space_after = Pt(6)
    run = p.add_run(text)
    size = {1: Pt(16), 2: Pt(14), 3: Pt(12), 4: Pt(10.5)}.get(level, Pt(10.5))
    set_run_font(run, name_cn='宋体', name_en='宋体', size=size, bold=True)

# ---------- 辅助函数：添加正文段落 ----------
def add_para(text, indent=True):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_before = Pt(0)
    pf.space_after = Pt(0)
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        pf.first_line_indent = Cm(0.74)  # 首行缩进2字符
    # 处理行内加粗 **text**
    parts = re.split(r'(\*\*.*?\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = p.add_run(part[2:-2])
            set_run_font(run, bold=True)
        else:
            run = p.add_run(part)
            set_run_font(run)
    return p

# ---------- 辅助函数：设置表格宽度 ----------
def set_table_width(table, width_cm=16.0):
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        tbl.insert(0, tblPr)
    tblW = tblPr.find(qn('w:tblW'))
    if tblW is None:
        tblW = parse_xml(f'<w:tblW {nsdecls("w")} w:w="{int(width_cm * 567)}" w:type="dxa"/>')
        tblPr.append(tblW)
    else:
        tblW.set(qn('w:w'), str(int(width_cm * 567)))

# ---------- 辅助函数：设置单元格边距 ----------
def set_cell_margins(cell, top=0, bottom=0, left=0.05, right=0.05):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{int(top * 20)}" w:type="dxa"/>'
        f'<w:bottom w:w="{int(bottom * 20)}" w:type="dxa"/>'
        f'<w:left w:w="{int(left * 100)}" w:type="dxa"/>'
        f'<w:right w:w="{int(right * 100)}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tcPr.append(tcMar)

# ---------- 辅助函数：表格转换 ----------
def add_table_from_markdown(table_lines, is_info_table=False):
    """将 markdown 表格文本转为 Word 表格，自动跳过分隔行"""
    rows_data = []
    for line in table_lines:
        line = line.strip()
        if not line or re.match(r'^\|[\|\-\=\:\s]+\|$', line):
            continue
        cells = [c.strip() for c in line.split('|')[1:-1]]
        rows_data.append(cells)

    if not rows_data:
        return

    num_cols = len(rows_data[0])
    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'

    if is_info_table:
        # 信息表：无边框，窄表格
        remove_table_borders(table)
        set_table_width(table, 12.0)
        for i, row_data in enumerate(rows_data):
            for j, cell_text in enumerate(row_data):
                cell = table.cell(i, j)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell, left=0.1, right=0.3)
                p = cell.paragraphs[0]
                p.clear()
                pf = p.paragraph_format
                pf.line_spacing = 1.5
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                run = p.add_run(cell_text)
                if j == 0:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    set_run_font(run, name_cn='宋体', size=Pt(12), bold=True)
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    set_run_font(run, name_cn='宋体', size=Pt(12))
    else:
        # 数据表：网格线，9pt 字体
        set_table_width(table, 16.0)
        for i, row_data in enumerate(rows_data):
            for j, cell_text in enumerate(row_data):
                cell = table.cell(i, j)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                set_cell_margins(cell)
                p = cell.paragraphs[0]
                p.clear()
                pf = p.paragraph_format
                pf.line_spacing = 1.0
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(cell_text)
                set_run_font(run, size=Pt(9))
                if i == 0:
                    run.font.bold = True

    doc.add_paragraph()  # 表后空一行

# ---------- 解析 Markdown 并生成 Word ----------
with open(MD_PATH, 'r', encoding='utf-8') as f:
    lines = f.readlines()

# 是否遇到了第一个数据表格（标记 info_table 之后）
past_info_table = False

i = 0
while i < len(lines):
    line = lines[i].rstrip()

    if not line:
        i += 1
        continue

    # 标题（支持 1-4 级）
    h_match = re.match(r'^(#{1,4})\s+(.+)$', line)
    if h_match:
        level = len(h_match.group(1))
        title = h_match.group(2)
        add_heading_styled(title, level)
        i += 1
        continue

    # 表格
    if line.startswith('|'):
        table_lines = []
        while i < len(lines) and lines[i].strip().startswith('|'):
            table_lines.append(lines[i].strip())
            i += 1
        has_separator = any(re.match(r'^\|[\|\-\=\:\s]+\|$', l) for l in table_lines)
        if has_separator and len(table_lines) >= 2:
            # 判断是否为信息表（表头为"项目"/"内容"）
            first_data_line = None
            for tl in table_lines:
                t = tl.strip()
                if not re.match(r'^\|[\|\-\=\:\s]+\|$', t):
                    first_data_line = t
                    break
            is_info = False
            if not past_info_table and first_data_line:
                cells = [c.strip() for c in first_data_line.split('|')[1:-1]]
                if len(cells) >= 2 and ('项目' in cells[0] or '课程' in cells[0]):
                    is_info = True
            add_table_from_markdown(table_lines, is_info_table=is_info)
            if is_info:
                past_info_table = True
        else:
            for tl in table_lines:
                add_para(tl.strip('|'))
        continue

    # 水平线
    if line.startswith('---') and len(line) >= 3:
        doc.add_paragraph()
        i += 1
        continue

    # 图片
    img_match = re.match(r'!\[(.*?)\]\((.*?)\)', line)
    if img_match:
        img_name = img_match.group(2)
        img_alt = img_match.group(1).strip() or '图'
        img_path_full = os.path.join(IMG_DIR, img_name)
        if os.path.exists(img_path_full):
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(img_path_full, width=Inches(6.0))
            # 图标题
            cap = doc.add_paragraph()
            cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            cap.paragraph_format.first_line_indent = None
            cap.paragraph_format.space_before = Pt(3)
            cap.paragraph_format.space_after = Pt(6)
            cr = cap.add_run(img_alt)
            set_run_font(cr, size=Pt(10.5))
        else:
            add_para(f'[图片未找到: {img_name}]')
        i += 1
        continue

    # 普通段落
    add_para(line)
    i += 1

# ---------- 保存 ----------
doc.save(OUT_PATH)
print(f"Word 文档已保存至：{OUT_PATH}")
